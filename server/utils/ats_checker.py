import os
import re
from PyPDF2 import PdfReader
from reportlab.lib.pagesizes import letter
from reportlab.lib.units import inch
import docx
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from io import BytesIO

# ======================================================
# RESUME VALIDATION (NEW!)
# ======================================================

def validate_resume_content(file_path):
    """
    Validate that uploaded file is actually a resume.
    Returns: (is_valid, error_message)
    """
    
    from utils.extract_text import extract_text
    
    try:
        text = extract_text(file_path)
        
        if not text or len(text.strip()) < 100:
            return False, "File appears to be empty or too short to be a resume"
        
        text_lower = text.lower()
        
        # 1. Check for resume indicators
        resume_keywords = [
            'experience', 'education', 'skills', 'work', 'job',
            'degree', 'university', 'college', 'bachelor', 'master',
            'project', 'internship', 'certificate', 'course'
        ]
        
        keyword_count = sum(1 for keyword in resume_keywords if keyword in text_lower)
        
        if keyword_count < 3:
            return False, "This doesn't appear to be a resume. Please upload a valid resume file."
        
        # 2. Check for contact information (email or phone)
        has_email = bool(re.search(r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b', text))
        has_phone = bool(re.search(r'(\+?\d{1,3}[-.\s]?)?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}', text))
        
        if not (has_email or has_phone):
            return False, "No contact information found. This doesn't look like a resume."
        
        # 3. Reject if it looks like a job description
        jd_keywords = [
            'job description', 'responsibilities:', 'requirements:', 
            'we are looking for', 'the ideal candidate', 'apply now',
            'job posting', 'vacancy', 'hiring for'
        ]
        
        jd_count = sum(1 for keyword in jd_keywords if keyword in text_lower)
        if jd_count >= 2:
            return False, "This appears to be a job description, not a resume. Please upload your resume."
        
        # 4. Reject if it looks like ID card/document
        id_keywords = ['pan card', 'aadhar', 'passport', 'driving license', 'voter id', 'issued by']
        id_count = sum(1 for keyword in id_keywords if keyword in text_lower)
        
        if id_count >= 2:
            return False, "This appears to be an ID document, not a resume. Please upload your resume."
        
        # 5. Check minimum word count
        word_count = len(text.split())
        if word_count < 100:
            return False, f"Resume too short ({word_count} words). A typical resume has 300-800 words."
        
        return True, "Valid resume"
        
    except Exception as e:
        return False, f"Error validating file: {str(e)}"


# ======================================================
# ATS ISSUE DETECTION (IMPROVED)
# ======================================================

def detect_ats_issues(file_path):
    """
    Analyze resume file and detect ATS-breaking issues.
    NOW WITH VALIDATION!
    """
    
    # 1. VALIDATE FIRST
    is_valid, error_msg = validate_resume_content(file_path)
    if not is_valid:
        return {
            'score': 0,
            'issues': [{
                'type': 'invalid_file',
                'severity': 'critical',
                'title': 'Invalid File',
                'description': error_msg,
                'fix': 'Please upload a valid resume (PDF or DOCX format)',
                'icon': '❌'
            }],
            'warnings': [],
            'passed': False,
            'file_format': os.path.splitext(file_path)[1].lower(),
            'is_valid_resume': False
        }
    
    # 2. PROCEED WITH ATS CHECK
    issues = []
    warnings = []
    score = 100
    
    ext = os.path.splitext(file_path)[1].lower()
    
    # Check file format specific issues
    if ext == '.pdf':
        pdf_issues = check_pdf_issues(file_path)
        issues.extend(pdf_issues)
        score -= len(pdf_issues) * 15
        
    elif ext in ['.docx', '.doc']:
        docx_issues = check_docx_issues(file_path)
        issues.extend(docx_issues)
        score -= len(docx_issues) * 15
    
    # Check content-based issues
    content_issues, content_warnings = check_content_issues(file_path)
    issues.extend(content_issues)
    warnings.extend(content_warnings)
    score -= len(content_issues) * 10
    score -= len(content_warnings) * 5
    
    # Ensure score is between 0-100
    score = max(0, min(100, score))
    
    return {
        'score': score,
        'issues': issues,
        'warnings': warnings,
        'passed': score >= 70,
        'file_format': ext,
        'is_valid_resume': True
    }


def check_pdf_issues(pdf_path):
    """Check PDF-specific ATS issues"""
    issues = []
    
    try:
        reader = PdfReader(pdf_path)
        
        # Check for scanned/image-only PDFs
        text_content = ""
        for page in reader.pages:
            text_content += page.extract_text()
        
        if len(text_content.strip()) < 100:
            issues.append({
                'type': 'scanned_pdf',
                'severity': 'critical',
                'title': 'Scanned PDF Detected',
                'description': 'Your PDF appears to be scanned or image-based. ATS cannot read text from images.',
                'fix': 'We will convert this to a text-based PDF that ATS can read.',
                'icon': '📄'
            })
        
        # Check for form fields
        if '/AcroForm' in reader.trailer.get('/Root', {}):
            issues.append({
                'type': 'form_fields',
                'severity': 'high',
                'title': 'PDF Form Fields Detected',
                'description': 'PDF forms are not ATS-friendly. Data in form fields may not be parsed.',
                'fix': 'We will convert form fields to plain text in your PDF.',
                'icon': '📋'
            })
        
    except Exception as e:
        print(f"PDF check error: {e}")
    
    return issues


def check_docx_issues(docx_path):
    """Check DOCX-specific ATS issues"""
    issues = []
    
    try:
        doc = Document(docx_path)
        
        # 1. Check for tables
        if len(doc.tables) > 0:
            issues.append({
                'type': 'tables',
                'severity': 'high',
                'title': 'Tables Detected',
                'description': f'Found {len(doc.tables)} table(s). 80% of ATS systems cannot parse tables correctly.',
                'fix': 'We will convert tables to plain text format while preserving your content.',
                'icon': '📊',
                'count': len(doc.tables)
            })
        
        # 2. Check for images
        image_count = 0
        for rel in doc.part.rels.values():
            if "image" in rel.target_ref:
                image_count += 1
        
        if image_count > 0:
            issues.append({
                'type': 'images',
                'severity': 'high',
                'title': 'Images/Graphics Detected',
                'description': f'Found {image_count} image(s). ATS cannot read images or extract text from them.',
                'fix': 'We will remove images but keep all your text content intact.',
                'icon': '🖼️',
                'count': image_count
            })
        
        # 3. Check for non-standard fonts
        non_standard_fonts = check_fonts(doc)
        if non_standard_fonts:
            issues.append({
                'type': 'fonts',
                'severity': 'medium',
                'title': 'Non-Standard Fonts',
                'description': f'Found fonts: {", ".join(non_standard_fonts[:3])}. These may cause parsing errors in some ATS.',
                'fix': 'We will use Arial font (most ATS-compatible) while keeping your layout.',
                'icon': '🎨',
                'fonts': non_standard_fonts
            })
        
        # 4. Check for multi-column layout
        sections_with_columns = 0
        for section in doc.sections:
            if hasattr(section, '_sectPr') and section._sectPr.xpath('.//w:cols[@w:num]'):
                cols = section._sectPr.xpath('.//w:cols[@w:num]')
                if cols and int(cols[0].get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}num', 1)) > 1:
                    sections_with_columns += 1
        
        if sections_with_columns > 0:
            issues.append({
                'type': 'columns',
                'severity': 'high',
                'title': 'Multi-Column Layout',
                'description': 'Document uses multiple columns. ATS reads left-to-right and may jumble content.',
                'fix': 'We will convert to single-column layout while maintaining readability.',
                'icon': '📐'
            })
        
    except Exception as e:
        print(f"DOCX check error: {e}")
    
    return issues


def check_fonts(doc):
    """Check for non-standard fonts"""
    standard_fonts = {'calibri', 'arial', 'times new roman', 'georgia', 'helvetica', 'verdana'}
    found_fonts = set()
    
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            if run.font.name:
                font_name = run.font.name.lower()
                if font_name not in standard_fonts:
                    found_fonts.add(run.font.name)
    
    return list(found_fonts)[:5]  # Return max 5 fonts


def check_content_issues(file_path):
    """Check content-based issues"""
    issues = []
    warnings = []
    
    from utils.extract_text import extract_text
    text = extract_text(file_path)
    
    if not text:
        return issues, warnings
    
    text_lower = text.lower()
    
    # 1. Check for standard section headers
    required_sections = ['experience', 'education', 'skills']
    missing_sections = []
    
    for section in required_sections:
        if section not in text_lower:
            missing_sections.append(section.title())
    
    if missing_sections:
        warnings.append({
            'type': 'missing_sections',
            'severity': 'low',
            'title': 'Missing Standard Sections',
            'description': f'Recommended sections not found: {", ".join(missing_sections)}',
            'fix': 'Consider adding these sections for better ATS parsing.',
            'icon': '📑'
        })
    
    # 2. Check contact info
    has_email = bool(re.search(r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b', text))
    has_phone = bool(re.search(r'(\+?\d{1,3}[-.\s]?)?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}', text))
    
    if not has_email:
        warnings.append({
            'type': 'no_email',
            'severity': 'low',
            'title': 'Email Not Clearly Visible',
            'description': 'Make sure your email is prominently displayed.',
            'fix': 'Add email address at the top of your resume.',
            'icon': '📧'
        })
    
    if not has_phone:
        warnings.append({
            'type': 'no_phone',
            'severity': 'low',
            'title': 'Phone Number Not Found',
            'description': 'Contact number not detected.',
            'fix': 'Add phone number for recruiter contact.',
            'icon': '📱'
        })
    
    # 3. Check length
    word_count = len(text.split())
    if word_count < 200:
        warnings.append({
            'type': 'too_short',
            'severity': 'low',
            'title': 'Resume Quite Brief',
            'description': f'Your resume has {word_count} words. Typical resumes are 400-800 words.',
            'fix': 'Consider adding more details about achievements and skills.',
            'icon': '📏'
        })
    elif word_count > 1200:
        warnings.append({
            'type': 'too_long',
            'severity': 'low',
            'title': 'Resume Quite Lengthy',
            'description': f'Your resume has {word_count} words. Keep it concise (1-2 pages ideal).',
            'fix': 'Focus on most relevant and recent experience.',
            'icon': '📄'
        })
    
    return issues, warnings


# ======================================================
# SMART AUTO-FIX - MAINTAINS ORIGINAL FORMAT! (NEW!)
# ======================================================

def auto_fix_resume(file_path, output_path=None):
    """
    Fix ATS issues while MAINTAINING ORIGINAL FILE FORMAT.
    PDF → Fixed PDF
    DOCX → Fixed DOCX
    """
    
    ext = os.path.splitext(file_path)[1].lower()
    
    # Generate output path if not provided
    if output_path is None:
        base_name = os.path.splitext(file_path)[0]
        output_path = f"{base_name}_ATS_Optimized{ext}"  # SAME EXTENSION!
    
    print(f"🔧 Fixing {ext.upper()} file: {file_path}")
    
    if ext in ['.docx', '.doc']:
        return fix_docx_preserve_style(file_path, output_path)
    elif ext == '.pdf':
        return fix_pdf_preserve_format(file_path, output_path)
    else:
        raise ValueError(f"Unsupported file format: {ext}")


def fix_docx_preserve_style(docx_path, output_path):
    """
    Fix DOCX while preserving MAXIMUM formatting.
    Only changes what's necessary for ATS.
    """
    
    doc = Document(docx_path)
    
    # 1. Remove images (ATS can't read them anyway)
    for rel in list(doc.part.rels.values()):
        if "image" in rel.target_ref:
            del doc.part.rels[rel.rId]
    
    # 2. Convert tables to formatted text (PRESERVING INFO)
    for table in doc.tables:
        table_content = []
        
        for row in table.rows:
            row_data = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    row_data.append(cell_text)
            
            if row_data:
                # Create well-formatted text from table
                table_content.append(' | '.join(row_data))
        
        if table_content:
            # Get table position
            table_element = table._element
            table_parent = table_element.getparent()
            table_index = list(table_parent).index(table_element)
            
            # Insert as paragraphs
            for i, line in enumerate(table_content):
                para = doc.add_paragraph(line)
                # Keep formatting similar to document
                for run in para.runs:
                    run.font.name = 'Arial'
                    run.font.size = Pt(11)
                table_parent.insert(table_index + i, para._element)
            
            # Remove table
            table_parent.remove(table_element)
    
    # 3. Standardize fonts ONLY (keep bold, italics, colors)
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            # Keep existing formatting but change font
            current_bold = run.font.bold
            current_italic = run.font.italic
            current_size = run.font.size
            current_color = run.font.color.rgb if run.font.color and run.font.color.rgb else None
            
            # Change to Arial (ATS-safe)
            run.font.name = 'Arial'
            
            # Restore formatting
            if current_bold is not None:
                run.font.bold = current_bold
            if current_italic is not None:
                run.font.italic = current_italic
            if current_size:
                run.font.size = current_size
            if current_color:
                run.font.color.rgb = current_color
    
    # 4. Keep margins and page size
    # (Don't change - preserve user's template)
    
    doc.save(output_path)
    print(f"✅ Fixed DOCX saved (style preserved): {output_path}")
    
    return output_path


def fix_pdf_preserve_format(pdf_path, output_path):
    """
    Fix PDF issues while maintaining PDF format.
    Extracts text and creates clean, ATS-friendly PDF.
    """
    
    from utils.extract_text import extract_text
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    from reportlab.lib.enums import TA_LEFT, TA_CENTER
    
    # Extract text from original PDF
    text = extract_text(pdf_path)
    
    if not text:
        raise ValueError("Could not extract text from PDF")
    
    print(f"📄 Extracted {len(text)} characters from PDF")
    
    # Create new PDF with clean formatting
    doc_template = SimpleDocTemplate(
        output_path,
        pagesize=letter,
        leftMargin=0.75*inch,
        rightMargin=0.75*inch,
        topMargin=0.75*inch,
        bottomMargin=0.75*inch
    )
    
    # Build styles
    styles = getSampleStyleSheet()
    
    # Custom styles for ATS
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=14,
        textColor='black',
        spaceAfter=12,
        alignment=TA_CENTER,
        fontName='Helvetica-Bold'
    )
    
    heading_style = ParagraphStyle(
        'CustomHeading',
        parent=styles['Heading2'],
        fontSize=12,
        textColor='black',
        spaceAfter=6,
        spaceBefore=12,
        fontName='Helvetica-Bold'
    )
    
    body_style = ParagraphStyle(
        'CustomBody',
        parent=styles['BodyText'],
        fontSize=11,
        textColor='black',
        spaceAfter=6,
        alignment=TA_LEFT,
        fontName='Helvetica'
    )
    
    # Build content
    story = []
    lines = text.split('\n')
    
    for line in lines:
        line = line.strip()
        if not line:
            story.append(Spacer(1, 6))
            continue
        
        # Detect if line is a header
        is_title = False
        is_heading = False
        
        # First line is usually name (title)
        if len(story) == 0 and len(line) < 50:
            is_title = True
        # All caps or ends with colon = heading
        elif line.isupper() or line.endswith(':') or (len(line.split()) <= 4 and len(line) < 40):
            is_heading = True
        
        # Apply appropriate style
        try:
            if is_title:
                para = Paragraph(line, title_style)
            elif is_heading:
                para = Paragraph(line, heading_style)
            else:
                para = Paragraph(line, body_style)
            
            story.append(para)
        except Exception as e:
            # If paragraph fails, add as plain text
            print(f"⚠️ Skipping line due to formatting: {e}")
            continue
    
    # Build PDF
    doc_template.build(story)
    
    print(f"✅ Fixed PDF saved (format preserved): {output_path}")
    
    return output_path


# ======================================================
# COMPARISON DATA
# ======================================================

def get_before_after_comparison(original_path, fixed_path):
    """Generate comparison data"""
    
    from utils.extract_text import extract_text
    
    original_issues = detect_ats_issues(original_path)
    original_text = extract_text(original_path)
    
    fixed_issues = detect_ats_issues(fixed_path)
    fixed_text = extract_text(fixed_path)
    
    improvements = []
    original_issue_types = {issue['type'] for issue in original_issues['issues']}
    fixed_issue_types = {issue['type'] for issue in fixed_issues['issues']}
    
    resolved = original_issue_types - fixed_issue_types
    
    for issue_type in resolved:
        improvements.append({
            'type': issue_type,
            'message': f'Fixed: {issue_type.replace("_", " ").title()}'
        })
    
    return {
        'before': {
            'text': original_text[:500] + '...' if len(original_text) > 500 else original_text,
            'score': original_issues['score'],
            'issues': original_issues['issues'],
            'warnings': original_issues['warnings']
        },
        'after': {
            'text': fixed_text[:500] + '...' if len(fixed_text) > 500 else fixed_text,
            'score': fixed_issues['score'],
            'issues': fixed_issues['issues'],
            'improvements': improvements
        },
        'score_improvement': fixed_issues['score'] - original_issues['score']
    }