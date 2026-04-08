# server/utils/preview_generator.py

import os
import base64
from io import BytesIO
from PIL import Image, ImageDraw, ImageFont
from docx import Document
from PyPDF2 import PdfReader
import fitz  # PyMuPDF for PDF rendering and coordinate detection


# ======================================================
# MAIN PREVIEW GENERATOR
# ======================================================

def generate_resume_preview_with_highlights(file_path, issues):
    """
    Generate preview images of resume with red highlights on issues.
    
    Args:
        file_path: Path to PDF or DOCX file
        issues: List of detected ATS issues
    
    Returns:
        List of base64 encoded images (one per page) with highlights
    """
    
    ext = os.path.splitext(file_path)[1].lower()
    
    if ext == '.pdf':
        return generate_pdf_preview(file_path, issues)
    elif ext in ['.docx', '.doc']:
        return generate_docx_preview(file_path, issues)
    else:
        raise ValueError(f"Unsupported format: {ext}")


# ======================================================
# PDF PREVIEW WITH HIGHLIGHTS (using PyMuPDF - no poppler needed!)
# ======================================================

def generate_pdf_preview(pdf_path, issues):
    """
    Convert PDF to images using PyMuPDF and draw red boxes on problem areas.
    No poppler dependency required!
    """
    
    try:
        print(f"🖼️ Generating preview with highlights: {os.path.basename(pdf_path)}")
        
        # Open PDF with PyMuPDF
        doc = fitz.open(pdf_path)
        
        # Detect issue locations in PDF
        issue_locations = detect_pdf_issue_locations(pdf_path, issues)
        
        # Generate images for each page
        highlighted_images = []
        
        for page_num in range(len(doc)):
            page = doc[page_num]
            
            # Render page to image (300 DPI for good quality)
            mat = fitz.Matrix(300/72, 300/72)  # Scale from 72 DPI to 300 DPI
            pix = page.get_pixmap(matrix=mat)
            
            # Convert PyMuPDF pixmap to PIL Image
            img_data = pix.tobytes("png")
            img = Image.open(BytesIO(img_data))
            
            # Create a copy to draw on
            img_with_highlights = img.copy()
            draw = ImageDraw.Draw(img_with_highlights, 'RGBA')
            
            # Draw red boxes for issues on this page
            page_issues = [loc for loc in issue_locations if loc['page'] == page_num]
            
            for issue_loc in page_issues:
                # Draw red rectangle
                x1, y1 = issue_loc['x1'], issue_loc['y1']
                x2, y2 = issue_loc['x2'], issue_loc['y2']
                
                # Red border
                draw.rectangle(
                    [x1, y1, x2, y2],
                    outline='red',
                    width=8
                )
                
                # Semi-transparent red overlay
                draw.rectangle(
                    [x1, y1, x2, y2],
                    fill=(255, 0, 0, 30)  # Red with 30/255 opacity
                )
                
                # Add issue label
                label = issue_loc['label']
                try:
                    font = ImageFont.truetype("arial.ttf", 24)
                except:
                    font = ImageFont.load_default()
                
                # Label background
                try:
                    label_bbox = draw.textbbox((x1, y1 - 35), label, font=font)
                    draw.rectangle(label_bbox, fill='red')
                except:
                    # Fallback for older PIL versions
                    pass
                
                draw.text((x1, y1 - 35), label, fill='white', font=font)
            
            # Convert to base64
            buffered = BytesIO()
            img_with_highlights.save(buffered, format="PNG", quality=95)
            img_base64 = base64.b64encode(buffered.getvalue()).decode('utf-8')
            
            highlighted_images.append({
                'page': page_num + 1,
                'image': f"data:image/png;base64,{img_base64}",
                'width': img_with_highlights.width,
                'height': img_with_highlights.height,
                'issues': page_issues
            })
        
        doc.close()
        
        print(f"✅ Generated {len(highlighted_images)} preview images")
        return highlighted_images
    
    except Exception as e:
        print(f"❌ PDF preview generation error: {e}")
        import traceback
        traceback.print_exc()
        return []


def detect_pdf_issue_locations(pdf_path, issues):
    """
    Detect pixel coordinates of issues in PDF using PyMuPDF.
    """
    
    locations = []
    
    try:
        doc = fitz.open(pdf_path)
        
        for page_num in range(len(doc)):
            page = doc[page_num]
            
            # Get page dimensions (for coordinate scaling)
            page_rect = page.rect
            page_width = page_rect.width
            page_height = page_rect.height
            
            # Scale factor for 300 DPI images
            scale = 300 / 72  # PDF is 72 DPI, images are 300 DPI
            
            # Detect tables
            if any(issue['type'] == 'tables' for issue in issues):
                try:
                    tables = page.find_tables()
                    if hasattr(tables, 'tables'):
                        for table_num, table in enumerate(tables.tables):
                            bbox = table.bbox
                            locations.append({
                                'page': page_num,
                                'type': 'table',
                                'label': '🔴 TABLE',
                                'x1': int(bbox[0] * scale),
                                'y1': int(bbox[1] * scale),
                                'x2': int(bbox[2] * scale),
                                'y2': int(bbox[3] * scale)
                            })
                except Exception as e:
                    print(f"⚠️ Table detection failed: {e}")
            
            # Detect images
            if any(issue['type'] == 'images' for issue in issues):
                image_list = page.get_images()
                for img_num, img in enumerate(image_list):
                    # Get image bounding box
                    xref = img[0]
                    try:
                        img_bbox = page.get_image_bbox(xref)
                        locations.append({
                            'page': page_num,
                            'type': 'image',
                            'label': '🔴 IMAGE',
                            'x1': int(img_bbox.x0 * scale),
                            'y1': int(img_bbox.y0 * scale),
                            'x2': int(img_bbox.x1 * scale),
                            'y2': int(img_bbox.y1 * scale)
                        })
                    except:
                        pass
            
            # Detect non-standard fonts
            if any(issue['type'] == 'fonts' for issue in issues):
                try:
                    blocks = page.get_text("dict")["blocks"]
                    for block in blocks:
                        if "lines" in block:
                            for line in block["lines"]:
                                for span in line["spans"]:
                                    font_name = span.get("font", "").lower()
                                    if font_name and not any(std in font_name for std in ['arial', 'helvetica', 'calibri', 'times']):
                                        bbox = span["bbox"]
                                        locations.append({
                                            'page': page_num,
                                            'type': 'font',
                                            'label': '🟡 FONT',
                                            'x1': int(bbox[0] * scale),
                                            'y1': int(bbox[1] * scale),
                                            'x2': int(bbox[2] * scale),
                                            'y2': int(bbox[3] * scale)
                                        })
                                        break  # Only highlight first occurrence per line
                except Exception as e:
                    print(f"⚠️ Font detection failed: {e}")
        
        doc.close()
        
    except Exception as e:
        print(f"⚠️ Could not detect precise locations: {e}")
        # Fallback: return general page-level highlights
        locations = create_fallback_locations(pdf_path, issues)
    
    return locations


def create_fallback_locations(pdf_path, issues):
    """
    Fallback method: place highlights at top of pages if precise detection fails.
    """
    
    locations = []
    
    try:
        reader = PdfReader(pdf_path)
        
        for page_num in range(len(reader.pages)):
            y_offset = 100
            
            for issue in issues:
                if issue['type'] == 'tables':
                    locations.append({
                        'page': page_num,
                        'type': 'table',
                        'label': '🔴 TABLE DETECTED',
                        'x1': 100,
                        'y1': y_offset,
                        'x2': 800,
                        'y2': y_offset + 60
                    })
                    y_offset += 80
                
                elif issue['type'] == 'images':
                    locations.append({
                        'page': page_num,
                        'type': 'image',
                        'label': '🔴 IMAGE DETECTED',
                        'x1': 100,
                        'y1': y_offset,
                        'x2': 800,
                        'y2': y_offset + 60
                    })
                    y_offset += 80
    except Exception as e:
        print(f"⚠️ Fallback location creation failed: {e}")
    
    return locations


# ======================================================
# DOCX PREVIEW WITH HIGHLIGHTS
# ======================================================

def generate_docx_preview(docx_path, issues):
    """
    Generate text-based preview for DOCX with issue markers.
    (Full image conversion would require additional dependencies)
    """
    
    try:
        print(f"📄 Generating DOCX preview: {os.path.basename(docx_path)}")
        return generate_docx_text_preview(docx_path, issues)
    
    except Exception as e:
        print(f"❌ DOCX preview generation error: {e}")
        return []


def generate_docx_text_preview(docx_path, issues):
    """
    Generate text-based preview for DOCX with issue markers.
    """
    
    try:
        doc = Document(docx_path)
        
        preview_text = ""
        issue_markers = []
        
        para_count = 0
        
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if text:
                preview_text += text + "\n\n"
                para_count += 1
        
        # Check for tables
        if doc.tables and any(issue['type'] == 'tables' for issue in issues):
            issue_markers.append({
                'type': 'table',
                'message': f'🔴 {len(doc.tables)} TABLE(S) DETECTED - ATS cannot parse tables',
                'location': 'Multiple locations in document'
            })
        
        # Check for images
        try:
            image_count = sum(1 for rel in doc.part.rels.values() if "image" in rel.target_ref)
            if image_count > 0 and any(issue['type'] == 'images' for issue in issues):
                issue_markers.append({
                    'type': 'image',
                    'message': f'🔴 {image_count} IMAGE(S) DETECTED - ATS cannot read images',
                    'location': 'Multiple locations in document'
                })
        except Exception as e:
            print(f"⚠️ Could not detect images in DOCX: {e}")
        
        return [{
            'page': 1,
            'text': preview_text[:5000],  # Limit text length
            'markers': issue_markers,
            'type': 'text_preview'
        }]
    
    except Exception as e:
        print(f"❌ DOCX text preview error: {e}")
        return []


# ======================================================
# HELPER: Get Preview Dimensions
# ======================================================

def get_preview_info(file_path):
    """
    Get basic info about the file for preview generation.
    """
    
    ext = os.path.splitext(file_path)[1].lower()
    
    try:
        if ext == '.pdf':
            reader = PdfReader(file_path)
            return {
                'pages': len(reader.pages),
                'format': 'PDF'
            }
        
        elif ext in ['.docx', '.doc']:
            doc = Document(file_path)
            return {
                'pages': 1,  # DOCX doesn't have page concept until rendered
                'format': 'DOCX',
                'paragraphs': len(doc.paragraphs),
                'tables': len(doc.tables)
            }
    except Exception as e:
        print(f"⚠️ Could not get preview info: {e}")
    
    return None