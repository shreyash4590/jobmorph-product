# server/utils/extract_text.py

import os
import re
from PyPDF2 import PdfReader
from PyPDF2.errors import PdfReadError
import docx

# Try to import OCR libraries (Google Cloud Vision API)
try:
    from google.cloud import vision
    import fitz  # PyMuPDF
    OCR_AVAILABLE = True
    print("✅ OCR available via Google Cloud Vision API")
except ImportError as e:
    OCR_AVAILABLE = False
    print(f"⚠️ OCR not available: {e}")


# ======================================================
# CUSTOM EXCEPTIONS
# ======================================================

class ScannedPDFError(Exception):
    """Raised when a scanned PDF is detected but OCR is not available"""
    pass

class EncryptedPDFError(Exception):
    """Raised when a password-protected PDF is uploaded"""
    pass

class CorruptedFileError(Exception):
    """Raised when a file is corrupted or unreadable"""
    pass


# ======================================================
# CONSTANTS
# ======================================================

MAX_OCR_PAGES = 10          # Prevent Vision API cost overrun
MIN_TEXT_LENGTH = 50        # Minimum chars to consider extraction successful
MAX_FILE_SIZE_BYTES = 10 * 1024 * 1024  # 10MB hard limit


# ======================================================
# MAIN ENTRY
# ======================================================

def extract_text(file_path):
    """
    Extract text from PDF, DOCX, or TXT.

    Handles:
    - Text-based PDFs (fast extraction)
    - Scanned PDFs (OCR fallback if Vision API available)
    - Encrypted/password-protected PDFs (clear error)
    - Corrupted files (clear error)
    - DOCX including table text (fix for Issue #17)
    - TXT with encoding fallback

    Raises:
        FileNotFoundError  – file doesn't exist
        EncryptedPDFError  – password-protected PDF (T1.9, T2.5)
        CorruptedFileError – unreadable/corrupted file (T1.6)
        ScannedPDFError    – scanned PDF and OCR unavailable (T1.2)
        ValueError         – unsupported file type or empty file (T1.5, T1.8)
        RuntimeError       – extraction failed for other reasons
    """

    # ── File existence check ──────────────────────────────────────
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"File not found: {file_path}")

    # ── File size check ───────────────────────────────────────────
    # Fixes T1.7, T2.1 — catches >10MB before any processing
    file_size = os.path.getsize(file_path)
    if file_size == 0:
        raise ValueError(
            "The uploaded file is empty (0 bytes). "
            "Please upload a valid resume file."
        )
    if file_size > MAX_FILE_SIZE_BYTES:
        size_mb = file_size / (1024 * 1024)
        raise ValueError(
            f"File is too large ({size_mb:.1f}MB). "
            f"Maximum allowed size is 10MB. "
            f"Tip: Compress your PDF or remove embedded images to reduce size."
        )

    # ── Route by extension ────────────────────────────────────────
    ext = os.path.splitext(file_path)[1].lower()

    if ext == ".pdf":
        return extract_from_pdf_with_ocr(file_path)

    elif ext in {".docx", ".doc"}:
        return extract_from_docx(file_path)

    elif ext == ".txt":
        return extract_from_txt(file_path)

    else:
        raise ValueError(
            f"Unsupported file type: '{ext}'. "
            "Please upload a PDF, DOCX, or TXT file."
        )


# ======================================================
# PDF EXTRACTION (TEXT → OCR FALLBACK)
# ======================================================

def extract_from_pdf_with_ocr(pdf_path):
    """
    Step 1: Try normal text extraction (fast, free)
    Step 2: If empty → OCR fallback with Vision API

    Correctly distinguishes between:
    - Encrypted PDFs  → EncryptedPDFError
    - Corrupted PDFs  → CorruptedFileError
    - Scanned PDFs    → ScannedPDFError (or OCR if available)
    - Empty PDFs      → ValueError
    """

    # ── Step 1: Try text extraction ───────────────────────────────
    text, pdf_issue = extract_from_pdf_text(pdf_path)

    # If we got a specific error type, raise immediately — don't try OCR
    if pdf_issue == "encrypted":
        raise EncryptedPDFError(
            "This PDF is password-protected. "
            "ATS systems cannot read encrypted PDFs. "
            "Please remove the password protection and re-upload."
        )

    if pdf_issue == "corrupted":
        raise CorruptedFileError(
            "This PDF file appears to be corrupted or invalid. "
            "Please try re-saving your resume and uploading again."
        )

    # ── Good text extracted ───────────────────────────────────────
    if text.strip() and len(text.strip()) >= MIN_TEXT_LENGTH:
        print(f"✅ Text extracted from PDF ({len(text.strip())} chars)")
        return text.strip()

    # ── Step 2: OCR Fallback (scanned/image-only PDF) ────────────
    print("⚠️ No extractable text found. Attempting OCR with Vision API...")

    if not OCR_AVAILABLE:
        raise ScannedPDFError(
            "This PDF appears to be scanned or image-based — "
            "it contains no readable text. "
            "Please upload a text-based PDF (created digitally, not scanned). "
            "Tip: If you only have a scanned copy, try retyping it in Word and saving as PDF."
        )

    try:
        ocr_text = extract_from_pdf_ocr(pdf_path)

        if not ocr_text.strip() or len(ocr_text.strip()) < MIN_TEXT_LENGTH:
            raise ScannedPDFError(
                "Could not extract enough text from this PDF. "
                "The file may be blank, corrupted, or contain very low quality images."
            )

        print(f"✅ OCR successful: {len(ocr_text)} characters extracted")
        return ocr_text.strip()

    except ScannedPDFError:
        raise
    except Exception as e:
        print(f"❌ OCR failed: {e}")
        raise ScannedPDFError(
            f"OCR processing failed for this scanned PDF. "
            "Please upload a text-based PDF instead."
        )


def extract_from_pdf_text(pdf_path):
    """
    Extract text from text-based PDFs.

    Returns:
        (text: str, issue: str | None)
        issue is one of: None, "encrypted", "corrupted"

    Fixes:
        T1.6  – corrupted PDF returns clear error (not silent "")
        T1.9  – encrypted PDF returns EncryptedPDFError (not "scanned" message)
        T2.5  – same as T1.9
    """
    text = ""

    try:
        reader = PdfReader(pdf_path)

        # ── Encrypted PDF check ───────────────────────────────────
        # Must check BEFORE accessing pages
        if reader.is_encrypted:
            print("🔒 Encrypted PDF detected")
            return "", "encrypted"

        # ── Empty PDF check ───────────────────────────────────────
        if len(reader.pages) == 0:
            print("⚠️ PDF has 0 pages")
            return "", None

        # ── Extract text from all pages ───────────────────────────
        for i, page in enumerate(reader.pages):
            try:
                content = page.extract_text()
                if content:
                    text += content + "\n"
            except Exception as page_err:
                # Single page failure — skip it, continue with rest
                print(f"⚠️ Page {i+1} extraction failed: {page_err}")
                continue

    except PdfReadError as e:
        # PyPDF2-specific error = corrupted file
        print(f"❌ PDF corrupted: {e}")
        return "", "corrupted"

    except Exception as e:
        error_msg = str(e).lower()
        # Some encrypted PDFs throw generic errors instead of using is_encrypted
        if any(word in error_msg for word in ["encrypt", "password", "decrypt"]):
            print(f"🔒 Encryption-related error: {e}")
            return "", "encrypted"
        print(f"⚠️ PDF read error: {e}")
        return "", "corrupted"

    return text, None


def extract_from_pdf_ocr(pdf_path):
    """
    OCR fallback using Google Cloud Vision API.
    Works on GCP App Engine (no poppler/tesseract needed).

    Fixes:
        Issue #25 – added MAX_OCR_PAGES limit to prevent cost overrun
    """
    if not OCR_AVAILABLE:
        raise ImportError("Google Cloud Vision API not configured")

    text = ""

    try:
        client = vision.ImageAnnotatorClient()
        pdf_document = fitz.open(pdf_path)
        total_pages = len(pdf_document)

        # ── Page limit to prevent cost/timeout overrun ────────────
        pages_to_process = min(total_pages, MAX_OCR_PAGES)

        if total_pages > MAX_OCR_PAGES:
            print(f"⚠️ PDF has {total_pages} pages — processing first {MAX_OCR_PAGES} only")

        print(f"📄 Processing {pages_to_process}/{total_pages} page(s) with Vision API...")

        for page_num in range(pages_to_process):
            page = pdf_document[page_num]

            # 200 DPI = good quality, reasonable size for Vision API
            pix = page.get_pixmap(dpi=200)
            img_bytes = pix.tobytes("png")

            image = vision.Image(content=img_bytes)
            response = client.text_detection(image=image)

            # ── Vision API error handling ─────────────────────────
            if response.error.message:
                error_msg = response.error.message
                if "quota" in error_msg.lower():
                    raise Exception(
                        "OCR quota exceeded. Please try again later."
                    )
                elif any(w in error_msg.lower() for w in ["permission", "authentication", "auth"]):
                    raise Exception(
                        "OCR authentication failed. Please contact support."
                    )
                else:
                    raise Exception(f"OCR error on page {page_num + 1}: {error_msg}")

            # ── Extract text ──────────────────────────────────────
            if response.text_annotations:
                page_text = response.text_annotations[0].description
                text += page_text + "\n\n"
                print(f"✅ Page {page_num + 1}/{pages_to_process} — {len(page_text)} chars")
            else:
                print(f"⚠️ Page {page_num + 1}/{pages_to_process} — no text found")

        pdf_document.close()

        if not text.strip():
            raise Exception("No text found in any page after OCR")

        print(f"✅ Vision API OCR complete: {len(text)} chars from {pages_to_process} pages")

    except Exception as e:
        print(f"❌ OCR processing failed: {e}")
        raise

    return text


# ======================================================
# DOCX EXTRACTION
# ======================================================

def extract_from_docx(docx_path):
    """
    Extract text from DOCX including:
    - Regular paragraphs
    - Text inside tables (CRITICAL FIX — Issue #17)
    - Headers and footers

    Fix:
        Issue #17 — original code used only document.paragraphs
        which completely missed all table content (skills, education,
        experience sections often stored in tables).

    Test coverage:
        T1.3 — DOCX extraction now includes table text
    """
    try:
        document = docx.Document(docx_path)
        text_parts = []

        # ── 1. Paragraphs ─────────────────────────────────────────
        for p in document.paragraphs:
            if p.text.strip():
                text_parts.append(p.text.strip())

        # ── 2. Tables (previously missing — Fix #17) ─────────────
        for table in document.tables:
            for row in table.rows:
                row_texts = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        row_texts.append(cell_text)
                if row_texts:
                    # Join cells with separator so context is preserved
                    text_parts.append(" | ".join(row_texts))

        # ── 3. Headers and Footers ────────────────────────────────
        for section in document.sections:
            for header_footer in [section.header, section.footer]:
                try:
                    for p in header_footer.paragraphs:
                        if p.text.strip():
                            text_parts.append(p.text.strip())
                except Exception:
                    # Header/footer access can fail on some DOCX files — skip
                    pass

        # ── Validate result ───────────────────────────────────────
        if not text_parts:
            raise ValueError(
                "DOCX file appears to be empty or contains no readable text. "
                "Please ensure your resume has actual text content."
            )

        combined = "\n".join(text_parts).strip()

        if len(combined) < MIN_TEXT_LENGTH:
            raise ValueError(
                f"DOCX file contains very little text ({len(combined)} characters). "
                "Please upload a complete resume."
            )

        print(f"✅ DOCX extracted: {len(combined)} chars "
              f"({len(document.paragraphs)} paragraphs, "
              f"{len(document.tables)} tables)")

        return combined

    except (ValueError, RuntimeError):
        raise
    except Exception as e:
        raise RuntimeError(
            f"Could not read DOCX file: {str(e)}. "
            "Please ensure the file is a valid Word document (.docx)."
        )


# ======================================================
# TXT EXTRACTION
# ======================================================

def extract_from_txt(txt_path):
    """
    Extract text from TXT files.
    Tries UTF-8 first, falls back to latin-1 for older files.

    Test coverage:
        T1.4 — TXT extraction
        T2.4 — Special characters (emoji, Chinese, accented)
    """
    # ── Try UTF-8 first ───────────────────────────────────────────
    try:
        with open(txt_path, "r", encoding="utf-8", errors="ignore") as f:
            content = f.read().strip()
    except Exception:
        # ── Fallback to latin-1 ───────────────────────────────────
        try:
            with open(txt_path, "r", encoding="latin-1", errors="ignore") as f:
                content = f.read().strip()
        except Exception as e:
            raise RuntimeError(
                f"Could not read TXT file: {str(e)}. "
                "Please ensure the file is a valid text file."
            )

    # ── Validate content ──────────────────────────────────────────
    if not content:
        raise ValueError(
            "TXT file is empty. Please upload a file with content."
        )

    # Remove null bytes and normalize whitespace
    content = content.replace('\x00', '')
    content = re.sub(r'\n{3,}', '\n\n', content)  # Max 2 consecutive newlines

    if len(content) < MIN_TEXT_LENGTH:
        raise ValueError(
            f"TXT file contains very little text ({len(content)} characters). "
            "Please upload a complete resume."
        )

    print(f"✅ TXT extracted: {len(content)} chars")
    return content



# # server/utils/extract_text.py

# import os
# from PyPDF2 import PdfReader
# import docx

# # Try to import OCR libraries (Google Cloud Vision API)
# try:
#     from google.cloud import vision
#     import fitz  # PyMuPDF
#     OCR_AVAILABLE = True
#     print("✅ OCR available via Google Cloud Vision API")
# except ImportError as e:
#     OCR_AVAILABLE = False
#     print(f"⚠️ OCR not available: {e}")


# class ScannedPDFError(Exception):
#     """Raised when a scanned PDF is detected but OCR is not available"""
#     pass


# # ======================================================
# # MAIN ENTRY
# # ======================================================
# def extract_text(file_path):
#     """
#     Extract text from PDF, DOCX, or TXT.
#     - Text-based PDFs: Fast extraction
#     - Scanned PDFs: OCR fallback (if available)
#     """

#     if not os.path.exists(file_path):
#         raise FileNotFoundError(f"❌ File not found: {file_path}")

#     ext = os.path.splitext(file_path)[1].lower()

#     if ext == ".pdf":
#         return extract_from_pdf_with_ocr(file_path)

#     elif ext in {".docx", ".doc"}:
#         return extract_from_docx(file_path)

#     elif ext == ".txt":
#         return extract_from_txt(file_path)

#     else:
#         raise ValueError(f"❌ Unsupported file type: {ext}")


# # ======================================================
# # PDF EXTRACTION (TEXT → OCR FALLBACK)
# # ======================================================
# def extract_from_pdf_with_ocr(pdf_path):
#     """
#     1️⃣ Try normal text extraction (FAST)
#     2️⃣ If empty → OCR fallback with Vision API
#     """

#     # Try text extraction first
#     text = extract_from_pdf_text(pdf_path)

#     if text.strip() and len(text.strip()) >= 50:
#         print("✅ Text extracted from PDF")
#         return text.strip()

#     # ---- OCR FALLBACK ----
#     print("⚠️ No extractable text found. Attempting OCR with Vision API...")

#     if not OCR_AVAILABLE:
#         raise ScannedPDFError(
#             "This PDF appears to be scanned or image-based. "
#             "OCR service is not available on this server. "
#             "Please upload a text-based PDF or contact support."
#         )

#     try:
#         ocr_text = extract_from_pdf_ocr(pdf_path)
        
#         if not ocr_text.strip() or len(ocr_text.strip()) < 50:
#             raise ScannedPDFError(
#                 "Could not extract enough text from this PDF. "
#                 "The file may be blank, corrupted, or have very low quality images."
#             )
        
#         print(f"✅ OCR successful: {len(ocr_text)} characters extracted")
#         return ocr_text.strip()
        
#     except Exception as e:
#         if isinstance(e, ScannedPDFError):
#             raise
#         print(f"❌ OCR failed: {e}")
#         raise ScannedPDFError(
#             f"Failed to extract text from this scanned PDF: {str(e)}"
#         )


# def extract_from_pdf_text(pdf_path):
#     """
#     Extract text from text-based PDFs.
#     Image-only PDFs will return empty string.
#     """
#     text = ""
#     try:
#         reader = PdfReader(pdf_path)

#         if len(reader.pages) == 0:
#             return ""

#         for i, page in enumerate(reader.pages):
#             content = page.extract_text()
#             if content:
#                 text += content + "\n"

#     except Exception as e:
#         print(f"⚠️ PDF text extraction error: {e}")
#         return ""

#     return text


# def extract_from_pdf_ocr(pdf_path):
#     """
#     OCR fallback using Google Cloud Vision API
#     Works on App Engine Standard (no poppler/tesseract needed!)
#     """
#     if not OCR_AVAILABLE:
#         raise ImportError("Google Cloud Vision API not configured")
    
#     text = ""
#     try:
#         # Initialize Vision API client
#         client = vision.ImageAnnotatorClient()
        
#         # Open PDF with PyMuPDF (works without poppler!)
#         pdf_document = fitz.open(pdf_path)
#         total_pages = len(pdf_document)
        
#         print(f"📄 Processing {total_pages} page(s) with Vision API...")
        
#         for page_num in range(total_pages):
#             page = pdf_document[page_num]
            
#             # Convert page to image (PNG format)
#             # Using 200 DPI for good quality and reasonable processing time
#             pix = page.get_pixmap(dpi=200)
#             img_bytes = pix.tobytes("png")
            
#             # Send to Vision API
#             image = vision.Image(content=img_bytes)
#             response = client.text_detection(image=image)
            
#             # ✅ FIXED: Better error handling for Vision API
#             if response.error.message:
#                 error_msg = response.error.message
                
#                 # Provide helpful error based on error type
#                 if "quota" in error_msg.lower():
#                     raise Exception(
#                         "Vision API quota exceeded. Please try again later or contact support."
#                     )
#                 elif "permission" in error_msg.lower() or "authentication" in error_msg.lower():
#                     raise Exception(
#                         "Vision API authentication failed. Please contact support."
#                     )
#                 else:
#                     raise Exception(f"Vision API error: {error_msg}")
            
#             # Extract text from annotations
#             if response.text_annotations:
#                 # First annotation contains all text
#                 page_text = response.text_annotations[0].description
#                 text += page_text + "\n\n"
#                 print(f"✅ Page {page_num + 1}/{total_pages} processed ({len(page_text)} chars)")
#             else:
#                 print(f"⚠️ Page {page_num + 1}/{total_pages} has no text")
        
#         pdf_document.close()
        
#         if not text.strip():
#             raise Exception("No text found in any page")
        
#         print(f"✅ Vision API OCR complete: {len(text)} characters extracted")
        
#     except Exception as e:
#         print(f"❌ OCR processing failed: {e}")
#         raise

#     return text


# # ======================================================
# # DOCX EXTRACTION
# # ======================================================
# def extract_from_docx(docx_path):
#     try:
#         document = docx.Document(docx_path)
#         paragraphs = [p.text for p in document.paragraphs if p.text.strip()]
        
#         if not paragraphs:
#             raise ValueError("DOCX file appears to be empty")
        
#         return "\n".join(paragraphs).strip()
        
#     except Exception as e:
#         raise RuntimeError(f"DOCX extraction failed: {e}")


# # ======================================================
# # TXT EXTRACTION
# # ======================================================
# def extract_from_txt(txt_path):
#     try:
#         with open(txt_path, "r", encoding="utf-8", errors="ignore") as f:
#             content = f.read().strip()
            
#         if not content:
#             raise ValueError("TXT file is empty")
            
#         return content
        
#     except Exception as e:
#         raise RuntimeError(f"TXT extraction failed: {e}")